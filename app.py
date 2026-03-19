import streamlit as st
import pandas as pd
from janome.tokenizer import Tokenizer
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import io
from docx import Document
from docx.shared import Inches, Pt, Mm

# --- 1. ページ設定 ---
st.set_page_config(page_title="Reflection List Export", layout="wide")
st.title("📋 振り返り一覧・ワードクラウド作成システム")

# --- 2. サイドバー設定 ---
st.sidebar.header("📋 行事・学習の設定")
event_name = st.sidebar.text_input("行事名・授業名", value="(例)理科校外学習")
event_date = st.sidebar.date_input("実施日")

t = Tokenizer()

# --- 3. 単語抽出エンジン (ワードクラウド用) ---
def extract_words(text):
    words = []
    # 感情語などの抽出条件（必要に応じて調整してください）
    target_feelings = ["面白い", "楽しい", "凄い", "すごい", "わかる", "驚く", "難しい", "疲れる", "つまらない", "嫌だ", "迷う"]
    if not text or pd.isna(text): return ""
    tokens = list(t.tokenize(str(text)))
    i = 0
    while i < len(tokens):
        token = tokens[i]
        part = token.part_of_speech.split(",")[0]
        base = token.base_form
        if part == "名詞" and len(base) >= 2:
            if base not in ["こと", "もの", "よう", "そう", "これ", "それ"]: words.append(base)
        elif base in target_feelings or part == "形容詞":
            word_to_add = base
            if i + 1 < len(tokens):
                next_t = tokens[i+1]
                if next_t.base_form in ["ない", "ぬ", "ん"]:
                    word_to_add = base + "ない"
                    i += 1
            if word_to_add not in ["ある", "する", "いる"]: words.append(word_to_add)
        i += 1
    return " ".join(words)

# --- 4. Word作成用関数 ---
def create_word(img_bytes, event, date, feedback_list):
    doc = Document()
    
    # 余白設定 (20mm)
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
    
    # タイトル
    doc.add_heading(f'{event} 振り返りまとめ', 0)
    doc.add_paragraph(f'出力日：{date.strftime("%Y年%m月%d日")}')
    
    # 1. ワードクラウド表示
    doc.add_heading('1. 全体傾向（ワードクラウド）', level=1)
    img_stream = io.BytesIO(img_bytes)
    doc.add_picture(img_stream, width=Inches(6.0))
    
    # 2. 感想一覧
    doc.add_heading('2. 個別回答一覧', level=1)
    
    # 見やすさのため表形式で出力
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = '振り返り内容'
    
    # 各列の幅調整（概算）
    hdr_cells[0].width = Mm(10)
    hdr_cells[1].width = Mm(150)

    for i, text in enumerate(feedback_list):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = str(text)
        # フォントサイズの調整
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# --- 5. メイン処理 ---
uploaded_file = st.file_uploader("ファイルをアップロードしてください (CSV/Excel)", type=["xlsx", "csv"])

if uploaded_file is not None:
    df = pd.read_csv(uploaded_file) if uploaded_file.name.endswith('.csv') else pd.read_excel(uploaded_file)
    target_col = st.selectbox("分析・出力対象の列を選択してください", df.columns)
    
    if st.button("ワードクラウドと一覧を作成する"):
        with st.spinner("処理中..."):
            # データの準備
            raw_feedback_list = df[target_col].dropna().astype(str).tolist()
            all_text_combined = " ".join(raw_feedback_list)
            wakati_text = extract_words(all_text_combined)
            
            if not wakati_text.strip():
                st.warning("有効な単語が抽出できませんでした。")
            else:
                # ワードクラウド作成
                # フォントパスは環境に合わせて変更してください（以下は一般的なLinux/Streamlit Cloud用）
                FONT_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
                try:
                    wc = WordCloud(font_path=FONT_PATH, background_color="white", width=1000, height=500).generate(wakati_text)
                except:
                    # フォントが見つからない場合のフォールバック（Windows等）
                    wc = WordCloud(background_color="white", width=1000, height=500).generate(wakati_text)

                img_buf = io.BytesIO()
                wc.to_image().save(img_buf, format='PNG')
                img_bytes = img_buf.getvalue()
                
                # 画面表示
                st.subheader("📊 ワードクラウド")
                st.image(img_bytes)
                
                st.subheader("📝 回答一覧（プレビュー）")
                st.dataframe(df[[target_col]])

                # Wordファイル作成
                word_file = create_word(img_bytes, event_name, event_date, raw_feedback_list)
                
                st.download_button(
                    label="📄 ワードクラウドと一覧をWordで保存",
                    data=word_file,
                    file_name=f"{event_name}_振り返りまとめ.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
